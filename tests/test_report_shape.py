# -*- coding: utf-8 -*-
"""The SHAPE of the results documents, as against the numbers in them.

`test_report.py` checks the arithmetic. Nothing checked the layout, and the
layout is not decoration here: the table went through eight revisions on 23
August -- transposed, an oracle column added, metrics regrouped by the model
that produces them, three background bands, sub-headers separating each head
from its bidder, and a second scoring population added to the funnel rows. Every
one of those was a decision about what a reader would compare with what.

A REGRESSION IN LAYOUT IS SILENT IN A WAY A REGRESSION IN ARITHMETIC IS NOT. If
a number goes wrong a contrast moves and something downstream disagrees. If the
three all-rows rows quietly stop being emitted -- which is exactly what
`results_report.py` does when `tier1_allrows.json` is absent, by design -- the
document still looks complete and still reads as the record. The reader is not
told that a row is missing, because a missing row leaves no gap.

So these tests assert the things a reader relies on and no other test would
notice: that the funnel rows carry the COMPARABLE figure rather than the one
`results.json` records, that no own-rows variant has crept back in, that the
three sections exist, that the two Tier-2 heads are separated, and that the
footnotes explaining all of it survived.

THE DOCX IS CHECKED TOO, WHEN IT CAN BE. python-docx lives in v1's environment,
not this one, so those tests skip rather than fail where it is absent. They are
worth having anyway: the docx is what Ken reads, the shading and the landscape
rotation happen only in the docx, and nothing else in the suite opens one.
"""
from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
REPO = ROOT.parent
TOOLS = ROOT / "tools"
RESULTS_MD = ROOT / "docs" / "V2.2_Results.md"

FUNNEL = ["click AUC", "install AUC", "payer AUC"]
SECTIONS = ["**Tier 1**", "**Tier 2, CLF**", "**Tier 2, AFT**"]


@pytest.fixture(scope="module")
def md():
    if not RESULTS_MD.exists():
        pytest.skip("no results document on this machine")
    return RESULTS_MD.read_text(encoding="utf-8")


@pytest.fixture(scope="module")
def rr():
    spec = importlib.util.spec_from_file_location("rr", TOOLS / "results_report.py")
    m = importlib.util.module_from_spec(spec)
    sys.path.insert(0, str(TOOLS))
    spec.loader.exec_module(m)
    return m


# ------------------------------------------------------------- the markdown

def test_one_funnel_row_each_and_no_own_rows_variant(md):
    """Both populations were printed for one day and that was wrong: every row
    in this table carries a C2-C1 column, and an own-rows contrast is largely a
    difference in scoring population rather than in data layer."""
    for h in FUNNEL:
        assert "| %s |" % h in md, "%s has no row at all" % h
        assert "%s, own rows" % h not in md, (
            "%s still has an own-rows row. It publishes a contrast that must "
            "not be compared, in a table whose every row is a comparison." % h)
        assert "%s, all rows" % h not in md, (
            "%s is labelled 'all rows'. There is only one figure now, so the "
            "qualifier is noise; the footnote says which population it is." % h)


def test_the_funnel_rows_carry_the_all_rows_figure(md):
    """The row must hold the COMPARABLE number, not the one results.json
    records. C1's click AUC is about 0.684 scored on every test row and about
    0.641 scored on the rows it won, so the two are far apart and easy to tell
    apart. If this ever reads 0.64 the accessor has been pointed back at the
    per-view figure and the contrast beside it is two thirds artifact."""
    for line in md.splitlines():
        if line.startswith("| click AUC |"):
            c1 = float(line.strip("|").split("|")[1].strip().replace("**", ""))
            assert c1 > 0.67, (
                "C1's click AUC reads %.3f. Scored on every test row it is "
                "about 0.684; 0.641 is the own-rows figure, which must not be "
                "the one reported." % c1)
            return
    raise AssertionError("no click AUC row found")


def test_the_three_sections_are_present_and_in_order(md):
    at = [md.find(s) for s in SECTIONS]
    assert all(i > 0 for i in at), dict(zip(SECTIONS, at))
    assert at == sorted(at), "the sections are out of order: %s" % at


def test_each_tier_2_section_separates_its_head_from_its_bidder(md):
    """Added because the classifier's profit once sat four rows from the
    classifier's AUC, with the price head's somewhere else again."""
    # BOLD ITALIC, not italic alone. Bolded 24 August so they read as
    # headings; the italic is what stops `shade_sections` treating them as
    # section headers and painting six bands where three belong.
    assert md.count("| ***the head*** |") == 2
    assert md.count("| ***the bidder*** |") == 2
    assert "| *the head* |" not in md, "reverted to italic-only"


def test_the_footnotes_that_explain_the_table_survived(md):
    for frag in ["Bold = best of C1", "Oracle = the ceiling",
                 "The funnel AUCs are scored on every test row",
                 "+0.0220 scored like for like",
                 "Price head sigma"]:
        assert frag in md, "the footnote %r is gone" % frag


def test_the_sigma_footnote_reports_both_values(md):
    """It used to name C3's residuals and print C1's number. Both views' values
    must appear, because they differ and that difference is queued as Q4."""
    i = md.find("**Price head sigma**")
    assert i > 0
    line = md[i:md.find("\n", i)]
    assert "0.899" in line and "1.000" in line, line
    assert "Q4" in line, "the footnote should point at the queued mismatch"


DROPPED = ["mce_win", "spend rmse_log", "value_captured"]


def _levels_table(md):
    """Just the per-view table, not the contrast tables below it.

    The document holds three tables and they answer different questions. This
    one reports LEVELS per view; the two below report CONTRASTS. A metric can
    rightly be absent from the first and present in the second, so a test about
    one must not read the others.
    """
    i = md.index("| Metric | C1 |")
    j = md.index("\n\n", i)
    return md[i:j]


def test_the_dropped_metrics_stay_dropped(md):
    """Removed 23 August 2026, each for its own reason.

    `mce_win` is the largest single-bin calibration gap, so it is one bin of ten
    and moves with whichever bin happened to be sparsest. Nothing in the study
    rests on it.

    `spend rmse_log` scores only the location of a predictive distribution whose
    whole shape `spend CRPS` already scores, and its NAME was the active harm: it
    sat in Tier 1 six rows above the price head's `rmse_log` in Tier 2, two
    unrelated quantities on different scales reading as a pair.

    `value_captured` IS `share of oracle`, times the oracle's own 0.959. The CLF
    bidder reads 0.232 and 0.241 at C1, and 0.232/0.959 is 0.242. Two rows, one
    quantity. The share survives because a reader can tell what 0.24 means and
    cannot tell what 0.23 means without also knowing the ceiling is not 1.

    All three remain in `results.json`, and `value_captured`'s C3-C1 contrast --
    negative with all ten seeds agreeing -- remains in the agreement tables,
    which is where contrasts belong. This asserts they are out of the per-view
    LEVELS table.
    """
    # SCOPED TO THE LEVELS TABLE, and it has to be. The contrast table further
    # down the same document reports `value_captured` per bidder per contrast
    # and SHOULD: that is the finding's home. Asserting over the whole file
    # would forbid the thing the next test requires.
    levels = _levels_table(md)
    for label in DROPPED:
        assert "| %s |" % label not in levels, \
            "%s is back in the per-view levels table" % label
    # the survivors, so a future edit cannot take the wrong one of a pair
    assert "| rmse_log |" in levels, "the price head's rmse_log was removed too"
    assert "| spend CRPS |" in levels, "the spend head lost its remaining figure"
    assert "| ece_win |" in levels, "ece_win went with mce_win; only mce was dropped"
    assert levels.count("| share of oracle |") == 2, \
        "both bidders must keep their share-of-oracle row"


def test_the_value_captured_contrast_is_still_recorded_somewhere():
    """Dropping it from the levels table must not drop the FINDING.

    C3-C1 on `value_captured` is negative with all ten seeds agreeing: the price
    bidder buys less under SSP. That belongs in the contrast record, and if the
    agreement table ever loses it too the finding has no home.
    """
    p = ROOT / "docs" / "gates" / "agreement_10M.md"
    if not p.exists():
        pytest.skip("no agreement table on this machine")
    body = p.read_text(encoding="utf-8")
    assert "value_captured" in body, \
        "the levels table dropped it, so the agreement table must still carry it"


def test_no_row_is_printed_as_nan(md):
    bad = [l for l in md.splitlines() if l.startswith("|") and "nan" in l.lower()]
    assert not bad, bad[:3]


def test_the_metric_list_and_the_document_agree(rr, md):
    """Every non-section METRICS label must appear as a row, so adding a metric
    and forgetting to regenerate is caught."""
    labels = [m[0] for m in rr.METRICS
              if m[1] is not None and not m[0].startswith("*")]
    missing = [l for l in labels if "| %s |" % l not in md]
    assert not missing, (
        "in METRICS but not in the document, so it is stale: %s" % missing)


# ------------------------------------------------------------------ the docx

def _docx(name):
    pytest.importorskip("docx", reason="python-docx lives in v1's environment")
    from docx import Document
    p = REPO / name
    if not p.exists():
        pytest.skip("%s has not been exported on this machine" % name)
    return Document(str(p))


def test_the_results_docx_matches_the_markdown_row_for_row():
    d = _docx("Training_Results v2.2.docx")
    labels = [r.cells[0].text.strip() for r in d.tables[0].rows]
    for h in FUNNEL:
        assert h in labels, "%s is missing from the Word table" % h
        assert not [l for l in labels if l.startswith("%s," % h)], (
            "%s has a qualified variant in the docx that the markdown does not "
            "have, so the export is stale" % h)
    assert labels.count("the head") == 2 and labels.count("the bidder") == 2


def test_the_results_docx_keeps_its_three_bands():
    """`--shade-sections` is opt-in, so a re-export that forgets the flag loses
    the banding and nothing else complains."""
    d = _docx("Training_Results v2.2.docx")     # skips first if python-docx is absent
    from docx.oxml.ns import qn
    fills = []
    for r in d.tables[0].rows:
        shd = r.cells[0]._tc.find(qn("w:tcPr"))
        f = shd.find(qn("w:shd")) if shd is not None else None
        fills.append(f.get(qn("w:fill")) if f is not None else None)
    bands = {f for f in fills if f and f != "D9D9D9"}
    assert len(bands) == 3, (
        "expected three section bands, found %d. Was the export run without "
        "--shade-sections? `tools/make_report.py` passes it." % len(bands))


def test_every_direction_in_METRICS_is_one_the_bolder_knows(rr):
    """The bug this catches bolded the WORST cell and looked fine.

    `_best` used to read `{"min":..., "near1":...}.get(direction, highest_wins)`,
    so any direction it did not recognise silently became "highest wins".
    `near0` was not in that dict, and `bias_log` is the row where zero is perfect
    -- so the table bolded +0.455 over -0.002 and told a reader the most biased
    view was the best one. A typo in any other direction string would have done
    the same, invisibly.
    """
    for label, acc, spec, direction, _ in rr.METRICS:
        if acc is None:
            continue
        assert direction in ("max", "min", "near1", "near0"), \
            "%s has direction %r, which _best does not know" % (label, direction)


def test_near0_bolds_the_value_closest_to_zero(rr):
    """The specific case, on the specific numbers that were wrong."""
    got = rr._best([0.455, 0.455, -0.002, -0.002], "near0", "%+.3f")
    assert got == {2, 3}, (
        "near0 must bold the two views nearest zero, got %s. If this is {0, 1} "
        "the fallback is back and bias_log is bolding its worst cell." % got)
    # and the other three directions still behave
    assert rr._best([1, 2, 3, 4], "max", "%.3f") == {3}
    assert rr._best([1, 2, 3, 4], "min", "%.3f") == {0}
    assert rr._best([0.5, 0.9, 1.1, 2.0], "near1", "%.3f") == {1}


def test_an_unknown_direction_raises_rather_than_guessing(rr):
    """Silence was the whole defect. A new metric with a mistyped direction must
    stop the run, not quietly get the max rule."""
    import pytest as _p
    with _p.raises(KeyError, match="unknown direction"):
        rr._best([1.0, 2.0], "nearest-to-pi", "%.3f")


# ------------------------------------------------- the AFT-only bidder set

AFT_MD = ROOT / "docs" / "V2.2_Results_AFT.md"


@pytest.fixture(scope="module")
def aft_md():
    if not AFT_MD.exists():
        pytest.skip("the AFT-only document has not been generated here")
    return AFT_MD.read_text(encoding="utf-8")


def test_the_aft_document_has_no_classifier_anywhere(aft_md):
    """Ken's decision of 24 August: the price head is the reported bidder.

    Checked over the WHOLE document, not just the levels table, because the
    contrast table further down lists a bidder per row and was the half that
    kept reporting `clf` after the levels table stopped.
    """
    assert "Tier 2, CLF" not in aft_md
    assert "`clf`" not in aft_md, "the contrast table still lists the classifier"
    assert "**Tier 2, AFT**" in aft_md, "the price head's section went too"


def test_the_section_filter_reads_sub_headers_correctly(rr):
    """The bug this catches emptied a section header and kept its rows.

    `**Tier 2, CLF**` and `***the head***` both start with two asterisks, so a
    filter testing only for that treats the SUB-header as a new section: it
    stops skipping at `***the head***` and the classifier's rows reappear
    orphaned under Tier 1. Bold-and-not-italic is the same rule the shader uses.
    """
    both = [m[0] for m in rr.metrics_for(("clf", "aft"))]
    aft = [m[0] for m in rr.metrics_for(("aft",))]
    assert "**Tier 2, CLF**" in both and "**Tier 2, CLF**" not in aft
    assert "**Tier 2, AFT**" in aft, "the wrong section was dropped"
    # the classifier's OWN rows must go with its header, and the price head's
    # identically-named rows must stay
    assert both.count("win AUC") == 2 and aft.count("win AUC") == 1
    assert both.count("profit, USD") == 2 and aft.count("profit, USD") == 1
    assert both.count("***the head***") == 2 and aft.count("***the head***") == 1
    # Tier 1 is untouched by either
    for m in ("click AUC", "install AUC", "payer AUC", "spend CRPS", "EV level"):
        assert m in both and m in aft, "%s should survive both sets" % m


def test_the_two_bidder_document_still_exists_and_still_has_both(md):
    """Dropping a bidder from the report must not delete the report that has
    both. It is cited, and a reader who finds one table where they remember two
    has to be able to see both."""
    assert "**Tier 2, CLF**" in md and "**Tier 2, AFT**" in md


def test_the_gate_files_keep_every_bidder():
    """The report narrows. The RECORD does not. `agreement_*.md` is where the
    contrasts live and the classifier's SSP null is a result in its own right."""
    p = ROOT / "docs" / "gates" / "agreement_10M.md"
    if not p.exists():
        pytest.skip("no agreement table on this machine")
    body = p.read_text(encoding="utf-8")
    assert "`clf`" in body, "the gate file dropped the classifier; only the report should"
    assert "`aft`" in body
