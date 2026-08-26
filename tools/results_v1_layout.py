# -*- coding: utf-8 -*-
"""The v2.2 results in v1's layout: two contrast tables, ticks and shading.

    python tools/results_v1_layout.py [--out <file.docx>]

Matches `Training_Results v10 + CI - V1.docx` so the two can be read side by
side, with ONE change Ken asked for: v1 put all four views in a single table
with one contrast column, which meant the MMP and SSP comparisons shared a row
and only one of them could be shown. Here they are two tables.

    Table 1   C1, C2, oracle, and the C2 - C1 contrast     the MMP layer
    Table 2   C1, C3, oracle, and the C3 - C1 contrast     the SSP layer

WRITTEN WITH python-docx RATHER THAN THROUGH MARKDOWN, because the two things
that make this layout readable cannot survive pandoc: the tick and cross glyphs
in the last two columns, and the per-row background fill. `results_report.py`
stays the markdown path and this is the presentation one; both read the same
`results.json` files and neither computes a statistic the other does not.

TWO COLUMNS THAT ARE NOT THE SAME QUESTION, and v1 ran them side by side for
exactly this reason:

    Improved   did the richer view move in the BENEFICIAL direction? Direction
               is per metric -- higher for an AUC, lower for an error, toward
               1.0 for a bias ratio -- so this cannot be read off a sign.
    Verdict    is that movement SUPPORTED? The paired interval must clear zero
               and every seed must agree.

A metric can improve without support, which is the amber case and the reason
both columns exist. Green fills a row only when both are ticks.

NO-CLAIM ROWS carry neither. `profit_per_1k_wins` and `wins` are reported
because a reader wants them and are not contrasts anybody is asserting, so they
show a dash and take no fill rather than a cross they did not earn.
"""
from __future__ import annotations

import argparse
import io
import json
import sys
from pathlib import Path

import numpy as np

sys.path.insert(0, str(Path(__file__).resolve().parent))
from results_report import paired as _paired      # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
RUNS = ROOT / "output" / "runs"

GREEN = "C6EFCE"        # both ticks
RED = "FFC7CE"          # improved but unsupported, or not improved
TICK, CROSS, DASH = "✅", "❌", "–"

# (label, per-view accessor, printf, direction, oracle, claim)
#   direction  max | min | near1 | none
#   oracle     None -> dash, "def1" -> 1.0, else a callable on one seed
#   claim      False -> reported, never judged
SECTIONS = [
    ("Tier 1 training", [
        ("EV level", lambda v, r: r[v]["ev"]["ratio"], "%.3f", "near1", "def1", True),
        ("auc_click", lambda v, r: r[v]["heads"]["click"]["auc"], "%.4f", "max", None, True),
        ("auc_install", lambda v, r: r[v]["heads"]["install"]["auc"], "%.4f", "max", None, True),
        ("auc_payer", lambda v, r: r[v]["heads"]["payer"]["auc"], "%.4f", "max", None, True),
        # NO auc_spend EXISTS. The head is E(spend | payer), a regression on log
        # spend, and AUC ranks a binary label -- there is no class here to rank.
        # CRPS ONLY. `rmse_log_spend` rode along beside it for one day and was
        # dropped on 23 August: CRPS scores the whole predictive distribution,
        # rmse_log scores only its location, and no claim in the study rests on
        # the second. It also collided by name with the price head's
        # `rmse_log` two sections below, which is a different quantity on a
        # different scale. v1 reported rmse_spend; v2.2 does not, and the two
        # documents are not comparable anyway.
        ("crps_spend, E(spend|payer)", lambda v, r: r[v]["heads"]["spend"]["crps"],
         "%.3f", "min", None, True),
    ]),
    ("Tier 2 training, CLF head (win classifier)", [
        ("auc_win (CLF)", lambda v, r: r[v]["heads"]["win"]["auc"], "%.4f", "max", None, True),
        # `mce_win` dropped 23 August, with `results_report.py`. One bin of
        # ten, moving with whichever bin was sparsest, and nothing rests on it.
        ("ece_win (CLF)",
         lambda v, r: r[v]["heads"]["win"]["ece_at_recommended"], "%.4f", "min", None, True),
    ]),
    ("Tier 2 training, AFT head (price model)", [
        ("price rmse_log", lambda v, r: r[v]["heads"]["price"]["rmse_log"],
         "%.4f", "min", None, True),
        ("price bias_log", lambda v, r: r[v]["heads"]["price"]["bias_log"],
         "%+.4f", "near0", None, True),
        ("auc_win (AFT)", lambda v, r: r[v]["heads"]["win_price"]["auc"],
         "%.4f", "max", None, True),
        ("ece_win (AFT)",
         lambda v, r: r[v]["heads"]["win_price"]["ece_at_recommended"],
         "%.4f", "min", None, True),
    ]),
    ("Tier 1+2, profit(bid) optimiser, CLF bidder", [
        # `value_captured` dropped 23 August, with results_report.py.
        # `value_vs_oracle` is the same number divided by the oracle's 0.959.
        ("share of oracle", lambda v, r: r[v]["economics"]["learned"]["value_vs_oracle"],
         "%.4f", "max", "def1", True),
        ("profit, total (USD)", lambda v, r: r[v]["economics"]["learned"]["profit"],
         "%,.2f", "max", lambda r: r["C1"]["economics"]["oracle"]["profit"], True),
        ("profit per 1k wins (USD)",
         lambda v, r: r[v]["economics"]["learned"]["profit_per_1k_wins"],
         "%,.2f", "none", None, False),
        ("n_won", lambda v, r: r[v]["economics"]["learned"]["wins"],
         "%,.0f", "none", lambda r: r["C1"]["economics"]["oracle"]["wins"], False),
    ]),
    ("Tier 1+2, profit(bid) optimiser, AFT bidder", [
        # dropped here too, see the CLF section above
        ("share of oracle", lambda v, r: r[v]["economics_price"]["learned"]["value_vs_oracle"],
         "%.4f", "max", "def1", True),
        ("profit, total (USD)", lambda v, r: r[v]["economics_price"]["learned"]["profit"],
         "%,.2f", "max", lambda r: r["C1"]["economics"]["oracle"]["profit"], True),
        ("profit per 1k wins (USD)",
         lambda v, r: r[v]["economics_price"]["learned"]["profit_per_1k_wins"],
         "%,.2f", "none", None, False),
        ("n_won", lambda v, r: r[v]["economics_price"]["learned"]["wins"],
         "%,.0f", "none", lambda r: r["C1"]["economics"]["oracle"]["wins"], False),
    ]),
]


# The two sections built on the win classifier. Dropped whole when the price
# head is reported alone, because a head and the bidder built on it are one
# story. Matched on the section title, which is what a reader sees.
CLF_SECTIONS = ("Tier 2 training, CLF head (win classifier)",
                "Tier 1+2, profit(bid) optimiser, CLF bidder")


def _exit(msg):
    print("\nREFUSING: %s\n" % msg)
    return 2


def sections_for(bidders, roas=None):
    secs = (SECTIONS if bidders == "both"
            else [(title, rows) for title, rows in SECTIONS
                  if title not in CLF_SECTIONS])
    if roas is None:
        return secs
    return [(title, [r for r in rows if r[0] not in ROAS_UNSCORABLE])
            for title, rows in secs]


# Rows the sweep cannot re-score, because it covers the BIDDER and not the
# heads. `ece_win` is scored on placed rows only, deliberately, so it genuinely
# moves with the target and a target-1.0 value would be a wrong number in a
# target-3.0 document. Dropped rather than shown stale.
ROAS_UNSCORABLE = ("ece_win (AFT)",)


def _with_ratio(rec, oracle_rec):
    """One sweep cell in results.json's policy-block shape.

    `value_vs_oracle` is not in the sweep and is not a new statistic: it is
    `value / oracle.value`, the same definition tools/backfill_value_vs_oracle.py
    uses. `mean_bid` comes out of the sweep as NaN and is read by nothing in this
    file, so it is passed through untouched rather than invented.
    """
    out = dict(rec)
    ceiling = oracle_rec.get("value")
    if ceiling:
        out["value_vs_oracle"] = rec["value"] / ceiling
    return out


def overlay_roas(rs, seeds, scale, target):
    """Replace the AFT bidder's economics with the sweep's, at `target`.

    WHY THIS EXISTS. The thirty results.json hold economics at a ROAS target of
    1.0 and re-scoring them in place needs about eight gigabytes a seed at 10M,
    which is what killed an earlier attempt at this change. The sweep already
    holds the same arithmetic at four targets over the same ten seeds, and its
    target-1.0 column reproduces each run's recorded economics exactly, which is
    the check that licenses the other three. So the document is built from the
    frozen runs for everything a target cannot move, and from the sweep for the
    bidder rows it can.

    THE RUNS ARE NOT TOUCHED. This overlays the dict in memory after it is
    loaded. results.json stays at 1.0 on disk, which keeps the generated record
    as it was generated and makes the re-scoring auditable rather than silent.

    CLASSIFIER ROWS CANNOT COME THROUGH HERE. The sweep scores the price head
    and the oracle only, so `--roas` refuses `--bidders both` in main() rather
    than filling the classifier's rows with the wrong bidder's numbers. The
    oracle IS re-scored, because it is gated by the same target.
    """
    sp = ROOT / "docs" / ("roas_sweep_%s.json" % scale)
    if not sp.exists():
        raise SystemExit("no %s.\n  Run: python tools/roas_sweep.py --scale %s"
                         % (sp, scale))
    j = json.loads(sp.read_text(encoding="utf-8"))
    tk = str(float(target))
    if tk not in {str(float(x)) for x in j["targets"]}:
        raise SystemExit("sweep %s holds targets %s, not %s"
                         % (sp.name, j["targets"], tk))
    for views, seed in zip(rs, seeds):
        cell = j["results"].get(str(seed))
        if cell is None:
            raise SystemExit("sweep has no seed %d at %s" % (seed, scale))
        # The sweep's oracle block carries fewer keys than a view's, because
        # tools/roas_sweep.py scores the oracle in its own pass. Both missing
        # keys are DERIVED here rather than left absent, so an accessor that
        # reads the oracle column does not hit a KeyError. `n` is the test-row
        # count, recovered from any view's own wins and win rate.
        orc = _with_ratio(cell["oracle"][tk], cell["oracle"][tk])
        ref = cell["C1"][tk]
        if ref.get("win_rate"):
            n_test = ref["wins"] / ref["win_rate"]
            orc.setdefault("win_rate", orc["wins"] / n_test)
        if orc.get("wins"):
            orc.setdefault("profit_per_1k_wins", 1000.0 * orc["profit"] / orc["wins"])
        for v in ("C1", "C2", "C3", "C4"):
            blk = views[v]
            # `truth_ev` is REMOVED rather than left at 1.0. Nothing in this
            # file reads it, and a stale number that raises nothing is worse
            # than a missing key that raises loudly.
            blk["economics_price"] = {"learned": _with_ratio(cell[v][tk], orc),
                                      "oracle": orc}
            blk["economics"] = dict(blk.get("economics") or {})
            blk["economics"]["oracle"] = orc
            blk["roas_target"] = float(target)
    return rs


def load(scale="10M", roas=None):
    paths = sorted((RUNS / scale).glob("seed*/results.json"))
    rs = [json.loads(p.read_text(encoding="utf-8"))["views"] for p in paths]
    if roas is None:
        return rs
    return overlay_roas(rs, [int(p.parent.name[4:]) for p in paths], scale, roas)


# THE INTERVAL IS IMPORTED, NEVER RESTATED. This file used to carry its own copy
# of the arithmetic under a docstring promising it was "the same arithmetic
# results_report.paired uses". It was, byte for byte, which is exactly why the
# copy was dangerous: nothing would have flagged the day it stopped being. A test
# forbids a second implementation (tests/test_report.py) and it was failing on
# this file. Importing keeps the two documents on one definition by construction.
def paired(vals):
    """`results_report.paired`, with this file's None-for-too-few convention."""
    st = _paired(vals)
    return None if st["n"] < 2 else st


def improved(base, rich, direction):
    """Did the richer view move the BENEFICIAL way? Not a sign test.

    Higher is better for an AUC and worse for an error, and a bias ratio is best
    at 1.0 rather than at either extreme, so the same signed difference means
    opposite things on different rows.
    """
    if direction == "max":
        return rich > base
    if direction == "min":
        return rich < base
    if direction == "near1":
        return abs(rich - 1.0) < abs(base - 1.0)
    if direction == "near0":
        return abs(rich) < abs(base)
    return None


def supported(st):
    """Interval clear of zero AND every seed agreeing."""
    if st is None or not np.isfinite(st["mean"]):
        return False
    return st["agree"] == st["n"] and (st["lo"] > 0 or st["hi"] < 0)


def structural_zero(diffs):
    """Every seed identical, so the layer cannot touch this metric at all.

    NOT A FAILED EFFECT, and calling it "not supported" in red would say the
    opposite of what it is. MMP adds funnel-label ROWS and no columns, and
    neither the win classifier nor the price head reads a funnel label, so
    C2 minus C1 on those rows is exactly zero on all ten seeds by construction.
    A reader should see that stated, not scored.
    """
    d = np.asarray([x for x in diffs if np.isfinite(x)], dtype=float)
    return len(d) > 0 and not d.any()


def fmt(x, spec):
    if x is None or not np.isfinite(x):
        return DASH
    if "," in spec:
        return format(x, spec.replace("%", "").replace("f", "f"))
    return spec % x


def build(doc, rs, rich, title, sections=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc.add_paragraph(title, style="Heading 2")
    head = ["Metric", "C1 (DSP)", "%s" % rich[1], "Oracle",
            "10M: mean [95%% CI], seeds agreeing" .replace("%%", "%"),
            "Improved", "Verdict"]
    n_rows = 1 + sum(1 + len(rows) for _, rows in SECTIONS)
    t = doc.add_table(rows=n_rows, cols=len(head))
    t.style = "Table Grid"

    for j, h in enumerate(head):
        c = t.cell(0, j)
        c.text = h
        for p in c.paragraphs:
            for r_ in p.runs:
                r_.bold = True
        shade(c, "D9D9D9")

    i = 1
    for section, rows in (SECTIONS if sections is None else sections):
        c = t.cell(i, 0)
        for j in range(1, len(head)):
            c = c.merge(t.cell(i, j))
        c.text = section
        for p in c.paragraphs:
            for r_ in p.runs:
                r_.bold = True
                r_.italic = True
        i += 1
        for label, acc, spec, direction, orc, claim in rows:
            g = lambda f: float(np.nanmean([f(r) for r in rs]))      # noqa: E731
            b_val = g(lambda r: acc("C1", r))
            r_val = g(lambda r: acc(rich[0], r))
            if orc is None:
                o_txt = DASH
            elif orc == "def1":
                o_txt = fmt(1.0, spec)
            else:
                o_txt = fmt(g(orc), spec)

            diffs = [acc(rich[0], r) - acc("C1", r) for r in rs] if claim else []
            st = paired(diffs) if claim else None
            if claim and structural_zero(diffs):
                ci = "0 on every seed: this layer cannot reach this metric"
                imp_txt, ver_txt, fill = DASH, "no effect by construction", None
            elif claim and st is not None:
                ci = "%s [%s, %s]  %d/%d" % (
                    fmt(st["mean"], "%+" + spec.lstrip("%")),
                    fmt(st["lo"], "%+" + spec.lstrip("%")),
                    fmt(st["hi"], "%+" + spec.lstrip("%")),
                    st["agree"], st["n"])
                imp = improved(b_val, r_val, direction)
                sup = supported(st)
                imp_txt = TICK if imp else CROSS
                # A SUPPORTED WORSENING GETS NO TICK. The claim being judged is
                # that the layer HELPS, so an effect that is solid and in the
                # wrong direction refutes it rather than supporting it. Ticking
                # it because the interval cleared zero would put a tick beside
                # a result that argues against the layer.
                if sup and imp:
                    ver_txt, fill = "%s  SUPPORTED" % TICK, GREEN
                elif sup and not imp:
                    ver_txt, fill = "%s  supported, AGAINST" % CROSS, RED
                else:
                    ver_txt, fill = "%s  not supported" % CROSS, RED
            else:
                ci, imp_txt, ver_txt, fill = DASH, DASH, "no claim", None

            for j, txt in enumerate([label, fmt(b_val, spec), fmt(r_val, spec),
                                     o_txt, ci, imp_txt, ver_txt]):
                cell = t.cell(i, j)
                cell.text = txt
                if fill:
                    shade(cell, fill)
                for p in cell.paragraphs:
                    if j:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j < 5 \
                            else WD_ALIGN_PARAGRAPH.CENTER
                    for r_ in p.runs:
                        r_.font.size = Pt(9)
            i += 1
    doc.add_paragraph()


def shade(cell, fill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    tcPr.append(el)


def main(argv=None):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--bidders", default="both", choices=["both", "aft"],
                    help="which bidders to report. `aft` drops the two "
                         "classifier sections and writes a separate file")
    ap.add_argument("--out", default=None)
    ap.add_argument("--roas", type=float, default=None,
                    help="re-score the bidder rows to this ROAS target from "
                         "docs/roas_sweep_<scale>.json. The runs on disk are "
                         "not touched. Requires --bidders aft, because the "
                         "sweep does not score the classifier")
    a = ap.parse_args(argv)
    if a.roas is not None and a.bidders == "both":
        return _exit(
            "--roas needs --bidders aft."
            "\n  The sweep scores the price head and the oracle only, so the "
            "classifier's rows cannot be re-scored, and would be left at the "
            "old target\n  beside re-scored ones in the same table.")

    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm

    rs = load("10M", roas=a.roas)
    if len(rs) < 2:
        print("REFUSING: %d seeds at 10M, an interval needs at least 2." % len(rs))
        return 2

    doc = Document()
    s = doc.sections[0]
    s.orient, s.page_width, s.page_height = WD_ORIENT.LANDSCAPE, s.page_height, s.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(s, m, Cm(1.5))

    doc.add_paragraph("T9Sim v2.2 — training results, 10M, n = 10 seeds"
                      + ("" if a.roas is None else ", ROAS target %g" % a.roas),
                      style="Title")
    doc.add_paragraph(
        "Generated by tools/results_v1_layout.py. Written, never edited. Laid out "
        "as v1's Training_Results v10 + CI so the two can be read side by side, "
        "with the four views split into two tables so the MMP and SSP contrasts "
        "each get their own column rather than sharing one.")
    if a.roas is not None:
        doc.add_paragraph(
            "The bidder rows are re-scored to a ROAS target of %g by "
            "tools/roas_sweep.py and are NOT read from the runs, which hold a "
            "target of 1.0. The ROAS gate is applied after the argmax, so a "
            "target changes which rows are bid on and never the bid on a placed "
            "row, and it cannot reach a fitted model. Every head row above the "
            "bidder therefore comes from the runs unchanged. The sweep's "
            "target-1.0 column reproduces each run's recorded economics exactly, "
            "wins and profit, in all four views on all ten seeds, and that is "
            "the check that licenses this one. ece_win is omitted because it is "
            "scored on placed rows only and the sweep cannot re-score it."
            % a.roas)
    doc.add_paragraph(
        "Improved asks whether the richer view moved the BENEFICIAL way, which is "
        "higher for an AUC, lower for an error and toward 1.0 for a bias ratio. "
        "Verdict asks whether that movement is supported: the paired interval "
        "must clear zero and all ten seeds must agree in direction. A row is "
        "green only when both are ticks. Rows marked no claim are reported "
        "because a reader wants them and are not contrasts being asserted.")

    secs = sections_for(a.bidders, roas=a.roas)
    build(doc, rs, ("C2", "C2 (+MMP)"),
          "Table 1. The MMP layer: C2 against C1", secs)
    build(doc, rs, ("C3", "C3 (+SSP)"),
          "Table 2. The SSP layer: C3 against C1", secs)

    # A SEPARATE FILE, never an overwrite. The two-bidder document is cited
    # and must stay quotable.
    default = ("Training_Results v2.2 + CI.docx" if a.bidders == "both"
               else "Training_Results v2.2 AFT + CI.docx")
    out = Path(a.out) if a.out else ROOT.parent / default
    # A LOCKED FILE IS A SENTENCE, NOT A STACK TRACE. Word holds an exclusive
    # handle on anything it has open, and the failure surfaced here as eighteen
    # lines of python-docx internals ending in PermissionError -- which reads
    # like the tool is broken rather than like the document is open. The other
    # writer, tools/safe_docx_export.py, has said this in one line since it was
    # built; this one had not caught up.
    try:
        doc.save(str(out))
    except PermissionError:
        raise SystemExit(
            "\nWord has %s open, so it cannot be written.\n"
            "  Close the document in Word and run this again.\n"
            "  Nothing else was changed." % out.name)
    print("wrote %s" % out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
